import os
import fitz  
import pytesseract
from PIL import Image
import re
import shutil
from datetime import datetime
from docx import Document 
from docx.shared import Inches, Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict 

# --- TESSERACT SMART PATHFINDER ---
tess_path = shutil.which("tesseract")
if tess_path:
    pytesseract.pytesseract.tesseract_cmd = tess_path
elif os.name == 'nt':  
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- 1. SETUP ---
INPUT_DIR = "Input"
OUTPUT_DIR = "Output"
PHOTOS_DIR = "Photos"
TEMPLATES_DIR = "Templates"

ZUA_FACILITIES = ["PGUM", "PGUA", "PGRO", "PGWT", "PGSN", "AJA", "UNZ", "UAM", "GRO", "SN"]

# --- 2. TEST TYPES ---
TEST_TYPES = {
    "ILS/L": "ILS Localizer testing",
    "ILS/G": "ILS Glideslope testing",
    "PROC/G": "RNAV Approach testing",
    "PROC/P": "RNAV Approach testing",
    "PROC/V": "RNAV Approach testing",
    "PROC/A": "RNAV Approach testing",
    "PROC/N": "Airway testing",
    "PROC/R": "Airway testing",
    "PROC/X": "Airway testing",
    "VTAC/V": "VOR + TAC Align Orbit",
    "VTAC/T": "VOR + TAC Radial testing",
    "APL/": "Runway Lighting testing",
    "NDBH/N": "NDB Check"
}

# --- 3. ADVANCED MANEUVER MAPPING ---
MANEUVERS_MAP = {
    "ILS Localizer testing": [
        {
            "name": "ILS 10 mile arc from localizer",
            "image": "ils_10_mile_arc.png",
            "details": [
                "Expected Time: Approximately 30-45 minutes.",
                "Arcs within 35 degrees on each side of the centerline.",
                "Conducted at 1500’ true altitude above field elevation.",
                "Speed is normally 170 knots IAS (Max 250).",
                "CRITICAL AREA: Objects in front of or overflying the localizer antenna will result in a repeat run. Ensure area remains clear.",
                "Expect 1 to 2 crossings for a LOC-only facility, or up to 4 successful crossings for a full ILS."
            ]
        }
    ],
    "ILS Glideslope testing": [
        {
            "name": "ILS holding pattern on localizer course",
            "image": "ils_holding_pattern.png",
            "details": [
                "Expected Time: Approximately 30-90 minutes (5-10 holding patterns).",
                "Flown along the localizer centerline from 6 NM to 2 NM outside the FAF.",
                "Conducted at 1500' true altitude (at or near glideslope intercept altitude).",
                "Speed is normally 170-200 knots IAS (Max 250).",
                "Maneuvering space: Left or right turnout can be dictated by ATC traffic needs."
            ]
        },
        {
            "name": "ILS low approach",
            "image": "ils_low_approach.png",
            "details": [
                "Expected Time: Approximately 30-45 minutes.",
                "Begins approximately 10 NM from the runway threshold, established on course and 500' above glide slope intercept.",
                "Aircraft will descend to 50' AGL and maintain 50' for the full length of the runway.",
                "Speed is normally 140-180 knots IAS.",
                "WAKE TURBULENCE: Avoidance is mandatory; the 50' run must be as stable as possible for signal recording.",
                "CRITICAL AREA: All runway and critical areas must be clear once the aircraft is established on course. Intervening objects will cause repeat runs."
            ]
        }
    ],
    "RNAV Approach testing": [
        {
            "name": "Required obstacle check (ROC) low approach",
            "image": "roc_low_approach.png",
            "details": [
                "MUST BE CONDUCTED IN VFR CONDITIONS: Aircraft will always accomplish this under VFR since the primary goal is visual obstacle identification.",
                "Normally flown from the FAF to the MAP in VMC conditions, maintaining 100’ below the published glidepath.",
                "Deviations left and right of course may be requested to visually check the height and location of suspect obstacles."
            ]
        },
        {
            "name": "General low approach",
            "image": "general_low_approach.png",
            "details": [
                "All segments flown to 100' below the lowest published minimums.",
                "Procedure survey information verified; this may result in up to 3 low approaches (including one opposite direction)."
            ]
        }
    ],
    "Airway testing": [
        {
            "name": "Airway",
            "image": "airway.png",
            "details": [
                "Generally flown in the direction of intended use at the published procedural altitude.",
                "During the evaluation, the flight check crew may request unexpected climbs or 360-degree turns."
            ]
        }
    ],
    "Runway Lighting testing": [
        {
            "name": "VGSI 2 Mile ARC from PAPI or VASI",
            "image": "vgsi_2_mile_arc.png",
            "details": [
                "Conducted 600' above field elevation to verify angular coverage and correct baffling on the VGSI.",
                "Arcs are similar to an ILS check but confined within 2 NM of the threshold.",
                "Recording is conducted from 10 degrees to 10 degrees on each side of the runway centerline.",
                "Aircraft positioning can be accommodated via left or right downwind to intercept the 2 NM arc.",
                "Expect a minimum of 2 arcs to be performed."
            ]
        },
        {
            "name": "Visual Glide slope indicator (VGSI) Low Approach",
            "image": "vgsi_low_approach.png",
            "details": [
                "Approximately 4 to 5 NM final, flown 1000-1500' above field elevation down to a 50' full-length low approach.",
                "Some approaches will be longer finals, below path, and left or right of course by 10 degrees to check obstacle clearance.",
                "Many aircraft will break off the approach prior to the threshold.",
                "Expect a minimum of 2 runs."
            ]
        }
    ],
    "VOR + TAC Align Orbit": [
        {
            "name": "Orbit",
            "image": "orbit.png",
            "details": [
                "Start point is flexible and can be based on ATC efficiency or traffic needs.",
                "Altitudes will vary depending on the distance of the orbit (typically 3,000 - 5,000' AGL).",
                "Speed is normally 170-250 knots IAS."
            ]
        }
    ],
    "VOR + TAC Radial testing": [
        {
            "name": "Radial",
            "image": "radial.png",
            "details": [
                "Radial flight from the facility, usually a 5-10 mile segment between 5 and 25 NM out.",
                "Can be flown inbound or outbound from the NAVAID.",
                "Altitudes will vary depending on the specific checkpoint.",
                "Speed is normally 170-200 knots IAS."
            ]
        },
        {
            "name": "VOR VDME VORTAC TACAN Low Approach",
            "image": "vor_tacan_low_approach.png",
            "details": [
                "Approaches are flown to 100’ below the published procedural minimums.",
                "Signal recording ends exactly at the missed approach point."
            ]
        }
    ],
    "NDB Check": [
        {
            "name": "Orbit",
            "image": "orbit.png",
            "details": [
                "360-degree orbit around the facility.",
                "Flown at the service volume distance (e.g., 50 NM) at 1500' true altitude above the facility."
            ]
        },
        {
            "name": "General low approach",
            "image": "general_low_approach.png",
            "details": [
                "Additional low approaches may be conducted."
            ]
        }
    ]
}

def translate_operation(line, facility):
    if "Arrive At" in line or "ROTM" in line or "FUTENMA" in line:
        return None
        
    test_desc = "Unknown Testing"
    for code, desc in TEST_TYPES.items():
        if code in line:
            test_desc = desc
            break
            
    if test_desc == "Unknown Testing":
        return None
        
    runway = ""
    rwy_match = re.search(r'(?<![\.\/])\b([0O][1-9]|[1-2][0-9]|3[0-6])[LRC]?\b(?![\.\/])', line)
    if rwy_match:
        clean_rwy = rwy_match.group(0).replace('O', '0').upper()
        runway = f"RWY {clean_rwy} "
        
    return f"{facility} {runway}{test_desc}", test_desc

# --- 4. MEMO GENERATOR ---
def create_overview_memo(translated_results):
    print("Generating official ZUA Overview Memo...")
    template_path = os.path.join(TEMPLATES_DIR, "ZUA Blank Memo.docx")
    
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Could not open template! Error: {e}")
        return

    today_date = datetime.today().strftime('%B %d, %Y')
    
    for paragraph in doc.paragraphs:
        if "Date:" in paragraph.text:
            paragraph.text = f"Date:\t{today_date}"
        elif "To:" in paragraph.text:
            paragraph.text = "To:\t\tOperations Floor"
        elif "From:" in paragraph.text:
            paragraph.text = "From:\t\tSupport Team"
        elif "Subject:" in paragraph.text:
            paragraph.text = "Subject:\tUpcoming Flight Check Operations Overview"

    doc.add_paragraph("")
    doc.add_paragraph("Flight check operations are scheduled to be conducted within the ZUA airspace on the following dates. Please review the summary of expected testing below.")
    doc.add_paragraph("")
    
    summary_title = doc.add_paragraph()
    summary_run = summary_title.add_run("Operations Summary by Date:")
    summary_run.bold = True
    
    current_date_header = ""
    for date, operation, _ in translated_results:
        if date != current_date_header:
            doc.add_paragraph("") 
            date_header = doc.add_paragraph()
            date_run = date_header.add_run(date)
            date_run.bold = True
            date_run.underline = True
            current_date_header = date
            
        doc.add_paragraph(f"  • {operation}")

    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run("Note: This schedule is intended as an overview of expected testing. Flight check itineraries remain tentative and are subject to change. The flight check crew may alter the sequence, modify maneuvers, or request additional testing while airborne based on operational requirements or weather conditions.")
    disclaimer_run.italic = True

    if translated_results:
        first_date_string = translated_results[0][0] 
        try:
            current_year = datetime.today().year
            date_only = first_date_string.split(",")[1].strip() 
            date_with_year = f"{date_only} {current_year}"
            parsed_date = datetime.strptime(date_with_year, "%d %B %Y")
            formatted_date = parsed_date.strftime("%m.%d.%y")
        except Exception:
            formatted_date = "Upcoming"
    else:
        formatted_date = "Upcoming"

    output_filename = f"Flight Check {formatted_date} Memo.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc.save(output_path)
    print(f"Success! Memo saved to: {output_path}")

# --- 5. IN-DEPTH BRIEFING GENERATOR ---
def create_floor_briefing(translated_results):
    print("Generating Multi-Maneuver Floor Briefing with Diagrams...")
    doc = Document()
    
    doc.add_heading('ZUA Flight Check - Floor Briefing', 0)
    doc.add_paragraph("_" * 50)
    
    grouped_data = defaultdict(lambda: defaultdict(list))
    test_order = [] 
    
    for date, operation, test_desc in translated_results:
        facility_rwy = operation.replace(test_desc, "").strip()
        grouped_data[test_desc][date].append(facility_rwy)
        
        if test_desc not in test_order:
            test_order.append(test_desc)

    for test_desc in test_order:
        doc.add_heading(test_desc, level=1)
        
        doc.add_paragraph().add_run("Scheduled Operations:").bold = True
        
        for date_str, locations in grouped_data[test_desc].items():
            doc.add_paragraph().add_run(date_str).bold = True
            
            num_cols = 3
            table = doc.add_table(rows=0, cols=num_cols)
            
            for i in range(0, len(locations), num_cols):
                row_locs = locations[i:i+num_cols]
                row_cells = table.add_row().cells
                for j, loc in enumerate(row_locs):
                    row_cells[j].text = f"  • {loc}"
        
        p_req = doc.add_paragraph()
        p_req.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_req = p_req.add_run("Maneuvers to Expect:")
        run_req.bold = True
        run_req.font.size = Pt(14)
        
        maneuvers = MANEUVERS_MAP.get(test_desc, [])
        
        for maneuver in maneuvers:
            doc.add_heading(maneuver["name"], level=3)
            
            image_path = os.path.join(PHOTOS_DIR, maneuver["image"])
            
            if maneuver["image"] in ["ils_10_mile_arc.png", "ils_holding_pattern.png"]:
                layout_table = doc.add_table(rows=1, cols=2)
                cell_left = layout_table.cell(0, 0)
                cell_right = layout_table.cell(0, 1)
                
                p_left = cell_left.paragraphs[0]
                if os.path.exists(image_path):
                    run_left = p_left.add_run()
                    run_left.add_picture(image_path, width=Inches(3.0)) 
                else:
                    p_left.add_run(f"[ Diagram missing: {maneuver['image']} ]")
                
                for i, detail in enumerate(maneuver["details"]):
                    if i == 0:
                        p_right = cell_right.paragraphs[0]
                        p_right.text = f"  • {detail}"
                    else:
                        cell_right.add_paragraph(f"  • {detail}")
            
            else:
                if os.path.exists(image_path):
                    if maneuver["image"] in ["vgsi_2_mile_arc.png", "vgsi_low_approach.png", "orbit.png", "radial.png", "vor_tacan_low_approach.png"]:
                        doc.add_picture(image_path, width=Inches(6.5))
                    else:
                        doc.add_picture(image_path, width=Inches(5.0))
                else:
                    doc.add_paragraph(f"[ Diagram missing: Please ensure {maneuver['image']} is in the Photos folder ]")
                
                for detail in maneuver["details"]:
                    doc.add_paragraph(f"  • {detail}")
            
        doc.add_paragraph("_" * 50)
        
    if translated_results:
        first_date_string = translated_results[0][0] 
        try:
            current_year = datetime.today().year
            date_only = first_date_string.split(",")[1].strip() 
            date_with_year = f"{date_only} {current_year}"
            parsed_date = datetime.strptime(date_with_year, "%d %B %Y")
            formatted_date = parsed_date.strftime("%m.%d.%y")
        except Exception:
            formatted_date = "Upcoming"
    else:
        formatted_date = "Upcoming"

    output_filename = f"Flight Check {formatted_date} Floor Briefing.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    
    doc.save(output_path)
    print(f"Success! Floor Briefing saved to: {output_path}")

# --- 6. FILTER & TRANSLATE ENGINE (STABLE PDF OCR) ---
def extract_filter_and_translate(pdf_filename):
    print(f"Opening {pdf_filename}...")
    file_path = os.path.join(INPUT_DIR, pdf_filename)
    doc = fitz.open(file_path)
    
    full_text = ""
    print(f"Reading all {doc.page_count} pages...")
    
    for i in range(doc.page_count):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        full_text += pytesseract.image_to_string(img) + "\n"
        
    translated_results = []
    current_date = "Unknown Date"
    
    for line in full_text.split('\n'):
        line = line.strip()
        
        if "Alternate Worklist" in line or "Task Remarks" in line:
            break
            
        date_match = re.match(r'^[A-Z][a-z]{2},\s\d{2}\s[A-Z][a-z]+', line)
        if date_match:
            current_date = date_match.group(0)
            continue
            
        for facility in ZUA_FACILITIES:
            if re.search(rf'\b{facility}\b', line):
                translation_data = translate_operation(line, facility)
                if translation_data: 
                    new_entry = (current_date, translation_data[0], translation_data[1])
                    
                    if new_entry not in translated_results:
                        translated_results.append(new_entry)
                break 
                
    create_overview_memo(translated_results)
    create_floor_briefing(translated_results)
